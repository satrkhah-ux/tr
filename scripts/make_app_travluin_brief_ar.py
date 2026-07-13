from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("output/documents/App-Travluin-Executive-Brief-AR.docx")


def set_rtl(paragraph, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    p_pr = paragraph._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")


def set_font(run, size=11, bold=False, color="1C3732"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_text(cell, text, bold=False, color="1C3732"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=10.5, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def para(doc, text="", bold=False, size=11, color="1C3732"):
    paragraph = doc.add_paragraph()
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return paragraph


def heading(doc, text, level=1):
    paragraph = doc.add_heading("", level=level)
    set_rtl(paragraph)
    run = paragraph.add_run(text)
    if level == 1:
        set_font(run, size=16, bold=True, color="185045")
    elif level == 2:
        set_font(run, size=13, bold=True, color="185045")
    else:
        set_font(run, size=12, bold=True, color="1F4D78")
    return paragraph


def bullet(doc, text):
    paragraph = doc.add_paragraph()
    set_rtl(paragraph)
    paragraph.paragraph_format.right_indent = Inches(0.15)
    run = paragraph.add_run("• " + text)
    set_font(run, size=10.8)


def callout(doc, title, text, fill="EAF7F1"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    paragraph = cell.paragraphs[0]
    set_rtl(paragraph)
    title_run = paragraph.add_run(title + "\n")
    set_font(title_run, size=12, bold=True, color="185045")
    body_run = paragraph.add_run(text)
    set_font(body_run, size=10.5, color="1C3732")
    doc.add_paragraph()


def make_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        shade_cell(table.cell(0, index), "EAF7F1")
        cell_text(table.cell(0, index), header, bold=True, color="185045")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cell_text(cells[index], value, bold=index == 0)
    doc.add_paragraph()


def build_document():
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("1C3732")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    para(doc, "ملف تعريفي للإدارة", bold=True, size=14, color="2AA87A")
    para(
        doc,
        "مشروع App Travluin\nاستنساخ وتطوير نظام إدارة شركة سياحية",
        bold=True,
        size=24,
        color="003C3A",
    )
    para(
        doc,
        "الغرض: عرض رؤية المشروع وخطة تحويل النظام الحالي إلى منصة تشغيل سياحية ذكية قابلة للتوسع.",
        size=11,
        color="557D78",
    )

    callout(
        doc,
        "الخلاصة التنفيذية",
        "يهدف المشروع إلى نقل النظام السياحي الحالي من نسخة تشغيلية تقليدية إلى منصة حديثة قابلة للتطوير، "
        "تجمع بين إدارة العملاء والموظفين والبكجات السياحية، وتوليد عروض احترافية بصيغة PDF أو روابط عميل قابلة للمشاركة، "
        "مع تمهيد واضح لإدخال الذكاء الاصطناعي ومحركات البحث الخارجية للفنادق والطيران.",
    )

    heading(doc, "1. تعريف الموقع / النظام الرئيسي", 1)
    para(
        doc,
        "النظام الرئيسي هو منصة إدارية داخلية مخصصة لشركة سياحية، تستخدم لإدارة العمليات اليومية المرتبطة بالعملاء، "
        "الموظفين، العروض السياحية، الفنادق، المدن، الدول، الطيران، الخدمات، التأشيرات، الشروط، والدليل التشغيلي.",
    )
    para(
        doc,
        "من خلال النظام الحالي يتم إنشاء عروض وبكجات سياحية للعملاء، ثم تحويلها إلى ملفات PDF احترافية أو إرسالها كبرامج سفر قابلة للمراجعة والمشاركة.",
    )

    heading(doc, "وظائف النظام الأصلي الأساسية", 2)
    for item in [
        "إدارة بيانات العملاء والأشخاص والشركات.",
        "إدارة الموظفين والصلاحيات والأدوار.",
        "إنشاء عروض وبكجات سياحية تشمل المدن والفنادق والطيران والتنقلات والخدمات.",
        "إدارة الدول والمدن والفنادق والرحلات والخدمات والتأشيرات.",
        "إنتاج ملفات PDF للعروض السياحية بصيغة منظمة وذات هوية بصرية.",
        "تجهيز برنامج العميل النهائي لإرساله ومراجعته قبل التأكيد.",
    ]:
        bullet(doc, item)

    heading(doc, "2. لماذا نقوم بالاستنساخ؟", 1)
    para(
        doc,
        "الاستنساخ ليس مجرد نسخ واجهة، بل هو مرحلة تأسيسية لفهم النظام الحالي وتحويله إلى بنية قابلة للتطوير السريع. "
        "الهدف هو امتلاك نسخة حديثة يمكن التحكم بها وتطويرها وإضافة مميزات جديدة عليها دون المساس بالنظام الأصلي أثناء العمل.",
    )
    heading(doc, "أهداف الاستنساخ", 2)
    for item in [
        "تحويل النظام الحالي إلى لغة وهيكلة يفهمها الذكاء الاصطناعي ويمكن تطويرها آليًا.",
        "إعادة بناء الواجهات والوظائف الأساسية بشكل مطابق قدر الإمكان للنظام الأصلي.",
        "فصل مرحلة التشغيل الحالية عن مرحلة الابتكار والتطوير المستقبلي.",
        "إنشاء قاعدة تقنية حديثة تسمح بإضافة API للفنادق والطيران والدفع والرسائل.",
        "تقليل الاعتماد على التدخل اليدوي في إنشاء العروض وتكرار البرامج.",
    ]:
        bullet(doc, item)

    heading(doc, "3. ماذا تم إنجازه حتى الآن؟", 1)
    for item in [
        "استنساخ نسبة كبيرة من واجهات النظام ولوحة الإدارة والقوائم الجانبية.",
        "إعادة بناء صفحات العملاء والموظفين والصلاحيات والبكجات والعروض.",
        "إنشاء صفحة إضافة عرض مشابهة لمراحل النظام الأصلي.",
        "إضافة مركز تشغيل ذكي يعرض رابط العميل وملف PDF وحالة العرض.",
        "إضافة رابط عميل يمكن فتحه خارج لوحة الإدارة لعرض البرنامج السياحي.",
        "إضافة مولد البكجات السياحية كتدفق عملي على مراحل.",
    ]:
        bullet(doc, item)

    heading(doc, "4. ماذا سنفعل في النسخة المطورة؟", 1)
    para(
        doc,
        "سنحوّل النسخة المستنسخة إلى منصة تشغيل سياحية حديثة باسم App Travluin، بحيث تصبح أداة إنتاج وتشغيل ومتابعة وليست مجرد قاعدة بيانات.",
    )

    make_table(
        doc,
        ["المكون", "الغرض", "القيمة للإدارة"],
        [
            ("مولد البكجات السياحية", "إنشاء برنامج العميل خطوة بخطوة من بيانات العميل حتى المعاينة", "تقليل الوقت والأخطاء وتوحيد جودة العروض"),
            ("مركز تشغيل العروض", "متابعة جاهزية العرض والرابط وملف PDF وحالة الإرسال", "رؤية تشغيلية فورية للإدارة والموظفين"),
            ("روابط العملاء", "صفحة عميل تفاعلية لكل عرض سياحي", "تجربة احترافية ومشاركة أسرع عبر واتساب أو البريد"),
            ("توليد PDF وصورة", "إخراج العرض بصيغ قابلة للإرسال والطباعة", "هوية موحدة ومخرجات قابلة للاعتماد"),
            ("ربط API للفنادق والطيران", "البحث عن الفنادق والأسعار والرحلات من مصادر خارجية", "رفع سرعة التسعير ودقة الاختيار"),
            ("ذكاء اصطناعي مساعد", "اقتراح برامج ومراجعة الأخطاء وتحسين النصوص", "رفع كفاءة الموظف وتقليل العمل اليدوي"),
        ],
    )

    heading(doc, "5. مولد البكجات السياحية المقترح", 1)
    para(
        doc,
        "مولد البكجات سيكون شاشة واحدة تعمل بمراحل ذكية، بحيث لا يرى الموظف كل المدخلات دفعة واحدة. "
        "يبدأ ببيانات العميل والرحلة، ثم يبحث النظام عن برامج محفوظة بنفس الدولة وعدد الأيام، ثم يسمح بالنسخ والتعديل أو إنشاء برنامج جديد.",
    )
    heading(doc, "تسلسل العمل داخل المولد", 2)
    for item in [
        "إدخال بيانات العميل: الاسم، الهاتف، الدولة، الوصول، المغادرة، عدد الأيام، عدد البالغين والأطفال حسب الأعمار.",
        "البحث عن عروض محفوظة بنفس الدولة وعدد الأيام، مع تقديم العروض المطابقة لعدد الأشخاص أولًا.",
        "نسخ برنامج سابق وبدء تعديله، أو إنشاء عرض مخصص من البداية.",
        "اختيار المدن وعدد الليالي والفنادق، ويظهر تطابق الليالي في هذه المرحلة فقط.",
        "اختيار الطيران والأسعار يدويًا أو مستقبلًا عبر API.",
        "اختيار الخدمات والشروط القابلة للتعديل.",
        "المعاينة النهائية ثم إنتاج PDF أو صورة أو ملف مكتوب أو رابط عميل.",
    ]:
        bullet(doc, item)

    heading(doc, "6. المخرجات التي سننتجها", 1)
    for item in [
        "نظام إدارة سياحي حديث باسم App Travluin.",
        "لوحة إدارة للموظفين والعملاء والبكجات والبيانات التشغيلية.",
        "مولد عروض وبكجات سياحية يعمل بمراحل سهلة.",
        "ملفات PDF احترافية بهوية Traveliun.",
        "روابط عميل تفاعلية لكل برنامج سياحي.",
        "إمكانية إنتاج صورة للعرض مناسبة للمشاركة السريعة.",
        "قاعدة بيانات منظمة للدول والمدن والفنادق والخدمات والشروط.",
        "قابلية ربط مستقبلية مع API للفنادق والطيران والرسائل والدفع.",
        "تمهيد لاستخدام الذكاء الاصطناعي في اقتراح البرامج وفحص الأخطاء وتحسين النصوص.",
    ]:
        bullet(doc, item)

    heading(doc, "7. التكاملات الخارجية المقترحة", 1)
    para(
        doc,
        "يمكن ربط النظام بمزودي خدمات خارجيين وفق توفر API رسمي أو اتفاقية شراكة. يفضّل البدء بمصادر رسمية مثل Amadeus للفنادق والطيران أو Hotelbeds للفنادق، "
        "مع تجنب أي ربط غير مصرح به عبر اسم مستخدم وكلمة مرور لمواقع لا توفر API.",
    )
    for item in [
        "API للفنادق: البحث حسب الدولة، المدينة، النجوم، نوع الفندق، السعر، والتوفر.",
        "API للطيران: البحث عن الرحلات والأسعار والأمتعة وأوقات الوصول والمغادرة.",
        "API للرسائل: إرسال رابط العرض عبر واتساب أو البريد أو SMS.",
        "API للدفع: تحويل العرض إلى فاتورة أو رابط دفع عند اعتماد العميل.",
        "ذكاء اصطناعي: اقتراح برنامج مبدئي بناءً على مدة الرحلة والدولة وعدد المسافرين.",
    ]:
        bullet(doc, item)

    heading(doc, "8. مراحل التنفيذ المقترحة", 1)
    make_table(
        doc,
        ["المرحلة", "الوصف", "المخرجات", "الأولوية"],
        [
            ("1", "استكمال مطابقة النظام الأصلي", "واجهات ووظائف مطابقة للنظام الحالي", "عالية"),
            ("2", "تطوير مولد البكجات", "تدفق إنشاء عروض منظم وسهل", "عالية"),
            ("3", "توليد PDF وروابط العملاء", "قوالب PDF وروابط مشاركة احترافية", "عالية"),
            ("4", "قاعدة بيانات تشغيلية", "دول، مدن، فنادق، خدمات، شروط", "متوسطة"),
            ("5", "التكاملات الخارجية", "API فنادق وطيران ورسائل ودفع", "متوسطة"),
            ("6", "الذكاء الاصطناعي", "اقتراح برامج وفحص جودة وتلخيص", "متقدمة"),
        ],
    )

    heading(doc, "9. القيمة المتوقعة للإدارة", 1)
    for item in [
        "رفع سرعة تجهيز العروض السياحية وتقليل الوقت المطلوب لكل برنامج.",
        "تحسين جودة مخرجات العملاء وتوحيد هوية الشركة في ملفات PDF والروابط.",
        "تقليل الأخطاء اليدوية في الليالي، الفنادق، الخدمات، والشروط.",
        "إتاحة متابعة أفضل لأداء الموظفين وحالة العروض.",
        "بناء نظام قابل للتوسع والربط مع مزودي الفنادق والطيران والدفع.",
        "تحويل الخبرة التشغيلية داخل الشركة إلى نظام ذكي قابل للتكرار والتطوير.",
    ]:
        bullet(doc, item)

    callout(
        doc,
        "التوصية",
        "اعتماد الاستنساخ الحالي كمرحلة تأسيسية، ثم الانتقال إلى تطوير App Travluin كنظام تشغيلي ذكي يبدأ بمولد البكجات وروابط العملاء وتوليد PDF، "
        "ثم يتوسع لاحقًا إلى API للفنادق والطيران والذكاء الاصطناعي.",
        "F4F8F6",
    )

    footer = doc.sections[0].footer.paragraphs[0]
    set_rtl(footer)
    run = footer.add_run("App Travluin - ملف تعريفي للإدارة")
    set_font(run, size=9, color="6F8F88")

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document().resolve())
